#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
translate_paper.py
------------------
一個把 PDF 論文翻譯成繁體中文的命令列工具，支援：
- 以 PyMuPDF（fitz）抽出可讀順序的文字
- 保留並跳過翻譯數學式（$...$、\[...\]、\(...\)、\begin{equation}...\end{equation} 等），翻譯後再還原
- 以段落切分 + 斷句分批翻譯，避免超過模型限制
- 多種翻譯後端（擇一安裝）：
  1) 本機 Hugging Face 模型（預設：facebook/m2m100_418M 或 Helsinki-NLP/opus-mt-en-zh）
  2) OpenAI API（需環境變數 OPENAI_API_KEY）
  3) DeepL API（需環境變數 DEEPL_API_KEY）
- 可選 OCR（針對掃描型 PDF；需安裝 tesseract）
- 產出 Markdown（.md）或 Docx（.docx）

使用範例：
  python translate_paper.py input.pdf --out output.md
  python translate_paper.py input.pdf --backend openai --openai-model gpt-4o-mini --out output.md
  python translate_paper.py input.pdf --backend deepl --out output.docx
  python translate_paper.py input.pdf --backend hf --hf-model facebook/m2m100_418M --src en --tgt zh-TW --out out.md

注意：
- 本工具預設把簡體轉繁體（台灣用語），若你本來就要簡體，可加上 --no-opencc。
- 若你的 PDF 非英語，建議指定 --hf-model 為多語模型（例如 facebook/m2m100_418M），或使用 OpenAI / DeepL 後端。
"""
from __future__ import annotations

import argparse
import re
import os
import sys
from dataclasses import dataclass
from typing import List, Tuple, Dict, Optional
import google.generativeai as genai

# ---------------------------
# 工具函式：偵測與處理數學式（mask/unmask）
# ---------------------------

MATH_PATTERNS = [
    (r"\$\$.*?\$\$", "BLOCK_DOLLAR"),   # $$ ... $$
    (r"\\\[.*?\\\]", "BLOCK_BRACKET"),  # \[ ... \]
    (r"\\begin\{equation\}.*?\\end\{equation\}", "ENV_EQUATION"),
    (r"\\begin\{align\*\}.*?\\end\{align\*\}", "ENV_ALIGNSTAR"),
    (r"\\begin\{align\}.*?\\end\{align\}", "ENV_ALIGN"),
    (r"\\begin\{eqnarray\}.*?\\end\{eqnarray\}", "ENV_EQNARRAY"),
    (r"\\\(.*?\\\)", "INLINE_PAREN"),   # \( ... \)
    (r"\$.*?\$", "INLINE_DOLLAR"),      # $ ... $
]

def mask_math(text: str) -> Tuple[str, Dict[str, str]]:
    """
    把數學式替換成唯一 placeholder，避免被翻譯器破壞。
    回傳 (masked_text, mapping)；最後記得 unmask。
    """
    mapping: Dict[str, str] = {}
    counter = 0
    def repl(match, tag):
        nonlocal counter
        counter += 1
        key = f"<<MATH_{tag}_{counter:05d}>>"
        mapping[key] = match.group(0)
        return key

    masked = text
    # 依序處理各種模式（先處理較長的 block）
    for pattern, tag in MATH_PATTERNS:
        masked = re.sub(pattern, lambda m, t=tag: repl(m, t), masked, flags=re.DOTALL)

    return masked, mapping

def unmask_math(text: str, mapping: Dict[str, str]) -> str:
    # 以鍵長度由長到短替換，避免子串相互影響
    for key in sorted(mapping.keys(), key=len, reverse=True):
        text = text.replace(key, mapping[key])
    return text

# ---------------------------
# PDF 文字抽取
# ---------------------------

def extract_paragraphs_from_pdf(pdf_path: str, use_ocr: bool=True) -> List[str]:
    """
    以 PyMuPDF 盡量依閱讀順序抽文字；若 use_ocr 啟用，會用 pdf2image + pytesseract。
    回傳段落列表（空段落會被略過）。
    """
    paragraphs: List[str] = []
    if use_ocr:
        try:
            from pdf2image import convert_from_path
            import pytesseract
        except Exception as e:
            print("使用 --ocr 需要安裝 pdf2image 與 pytesseract，並且系統需安裝 tesseract。", file=sys.stderr)
            raise

        pages = convert_from_path(pdf_path, dpi=300)
        for img in pages:
            raw = pytesseract.image_to_string(img, lang="eng")
            # 以雙換行分段
            parts = re.split(r"\n\s*\n", raw)
            paragraphs.extend([p.strip() for p in parts if p.strip()])
        return paragraphs

    # 非 OCR 路徑：PyMuPDF
    try:
        import fitz  # PyMuPDF
    except Exception:
        print("需要安裝 PyMuPDF（pymupdf）。請先 pip install pymupdf", file=sys.stderr)
        raise

    with fitz.open(pdf_path) as doc:
        for page in doc:
            # "blocks" 會保留相對合理的閱讀順序
            blocks = page.get_text("blocks")  # List[ (x0,y0,x1,y1, "text", block_no, block_type) ]
            blocks = sorted(blocks, key=lambda b: (round(b[1]), round(b[0])))
            for b in blocks:
                text = (b[4] or "").strip()
                # 跳過只含頁眉頁腳版次的短行
                if not text:
                    continue
                # 合併頁面中的 block，以雙換行斷段
                paragraphs.extend([p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()])
    return paragraphs

# ---------------------------
# 翻譯後端（介面 + 各實作）
# ---------------------------

@dataclass
class TranslateConfig:
    src_lang: str = "en"
    tgt_lang: str = "zh-TW"
    backend: str = "hf"  # hf | openai | deepl
    openai_model: str = "gpt-4o-mini"
    hf_model: Optional[str] = None  # "facebook/m2m100_418M" or "Helsinki-NLP/opus-mt-en-zh"
    use_opencc: bool = True
    opencc_config: str = "s2twp"  # 簡轉繁（台灣）


class TranslatorBackend:
    def translate_list(self, texts: List[str], cfg: TranslateConfig) -> List[str]:
        raise NotImplementedError


class HFTranslator(TranslatorBackend):
    """
    預設走 M2M100（可多語），若指定為 opus-mt-en-zh 則限英->中但較輕量。
    """
    def __init__(self, model_name: Optional[str]=None):
        self.model_name = model_name or "Helsinki-NLP/opus-mt-en-zh"
        self.tokenizer = None
        self.model = None

    def _ensure_loaded(self):
        if self.tokenizer is not None and self.model is not None:
            return
        from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
        self.model = AutoModelForSeq2SeqLM.from_pretrained(self.model_name)

    def translate_list(self, texts: List[str], cfg: TranslateConfig) -> List[str]:
        self._ensure_loaded()
        
        if "m2m100" in self.model_name:
            # M2M100 需要特殊處理
            tgt_lang = cfg.tgt_lang
            if tgt_lang in ["zh-TW", "zh-CN"]:
                tgt_lang = "zh"
            
            self.tokenizer.src_lang = cfg.src_lang
            results = []
            
            for text in texts:
                # 先檢查並截斷過長的文本
                tokens = self.tokenizer.encode(text, add_special_tokens=False)
                if len(tokens) > 400:  # 保留空間給特殊 tokens
                    truncated_tokens = tokens[:400]
                    text = self.tokenizer.decode(truncated_tokens, skip_special_tokens=True)
                
                # 編碼輸入，設定較小的 max_length
                inputs = self.tokenizer(text, return_tensors="pt", max_length=400, truncation=True, padding=True)
                
                # 移到同一設備
                if hasattr(self.model, 'device'):
                    inputs = {k: v.to(self.model.device) for k, v in inputs.items()}
                
                # 生成翻譯，調整參數
                generated_tokens = self.model.generate(
                    **inputs,
                    forced_bos_token_id=self.tokenizer.get_lang_id(tgt_lang),
                    max_new_tokens=400,  # 使用 max_new_tokens 而不是 max_length
                    num_beams=2,
                    early_stopping=True,
                    no_repeat_ngram_size=3,
                    repetition_penalty=1.2,
                    length_penalty=1.0,
                    pad_token_id=self.tokenizer.eos_token_id
                )
                
                # 解碼輸出
                translated = self.tokenizer.batch_decode(generated_tokens, skip_special_tokens=True)[0]
                results.append(translated)
            
            return results
        else:
            # 非 M2M100 模型的處理
            from transformers import pipeline
            
            # 使用較小的批次大小和更嚴格的長度限制
            results = []
            for text in texts:
                # 更嚴格的長度檢查
                tokens = self.tokenizer.encode(text, add_special_tokens=False)
                if len(tokens) > 300:  # 更保守的限制
                    truncated_tokens = tokens[:300]
                    text = self.tokenizer.decode(truncated_tokens, skip_special_tokens=True)
                
                # 每次只處理一個文本，避免批次問題
                pipe = pipeline(
                    "translation", 
                    model=self.model, 
                    tokenizer=self.tokenizer, 
                    device_map="auto",
                    max_length=400,
                    truncation=True
                )
                
                try:
                    gen = pipe(text, max_length=400, truncation=True)
                    if isinstance(gen, list) and len(gen) > 0:
                        results.append(gen[0]["translation_text"])
                    else:
                        results.append(text)
                except Exception as e:
                    print(f"翻譯失敗: {e}, 返回原文")
                    results.append(text)
            
            return results


class OpenAITranslator(TranslatorBackend):
    def __init__(self, model: str="gpt-4o-mini"):
        self.model = model
        # NOTE: 需要 `pip install openai>=1.0` 並設 OPENAI_API_KEY

    def translate_list(self, texts: List[str], cfg: TranslateConfig) -> List[str]:
        from openai import OpenAI
        import time
        client = OpenAI()
        out = []
        # 逐段翻譯，避免上下文過長；你也可以改成把多段塞在一個訊息裡
        for t in texts:
            prompt = (
                "你是一位嚴謹的論文翻譯助手。把以下英文學術段落翻成「精準、自然的繁體中文（台灣用語）」；"
                "保留引文標號與 DOI/URL，不要翻譯數學符號與變數名稱；必要時優化語序以更符合中文閱讀。\n\n"
                f"段落：\n{t}"
            )
            resp = client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a professional academic translator."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.2,
            )
            out.append(resp.choices[0].message.content.strip())
            time.sleep(0.1)  # 輕微節流
        return out


class DeepLTranslator(TranslatorBackend):
    def __init__(self):
        # 需要 pip install deepl 並設 DEEPL_API_KEY
        pass

    def translate_list(self, texts: List[str], cfg: TranslateConfig) -> List[str]:
        import deepl
        auth = os.environ.get("DEEPL_API_KEY")
        if not auth:
            raise RuntimeError("需要環境變數 DEEPL_API_KEY")
        translator = deepl.Translator(auth)
        target_lang = "ZH"  # DeepL: ZH = 中文，無法直接分繁/簡；可搭配 opencc 後處理
        out = []
        for t in texts:
            res = translator.translate_text(t, source_lang=cfg.src_lang.upper(), target_lang=target_lang)
            out.append(res.text)
        return out

# ---------------------------
# 主流程
# ---------------------------

def split_for_translation(paragraphs: List[str], max_chars: int=400) -> List[str]:  # 降低到 400
    """
    基於字數簡單分批，使用更小的 max_chars 避免超過模型限制。
    """
    batches: List[str] = []
    buf = []
    size = 0
    for p in paragraphs:
        if not p.strip():
            continue
        
        # 如果單個段落就很長，需要進一步切分
        if len(p) > max_chars:
            # 先把之前累積的加入 batches
            if buf:
                batches.append("\n\n".join(buf))
                buf = []
                size = 0
            
            # 將長段落按句子切分
            sentences = re.split(r'(?<=[.!?])\s+', p)
            current_chunk = ""
            for sent in sentences:
                if len(current_chunk) + len(sent) > max_chars and current_chunk:
                    batches.append(current_chunk.strip())
                    current_chunk = sent
                else:
                    current_chunk += " " + sent if current_chunk else sent
            if current_chunk:
                batches.append(current_chunk.strip())
        elif size + len(p) > max_chars and buf:
            batches.append("\n\n".join(buf))
            buf = [p]
            size = len(p)
        else:
            buf.append(p)
            size += len(p)
    
    if buf:
        batches.append("\n\n".join(buf))
    return batches
def build_backend(cfg: TranslateConfig) -> TranslatorBackend:
    if cfg.backend == "hf":
        return HFTranslator(model_name=cfg.hf_model)
    elif cfg.backend == "openai":
        return OpenAITranslator(model=cfg.openai_model)
    elif cfg.backend == "deepl":
        return DeepLTranslator()
    else:
        raise ValueError(f"未知後端：{cfg.backend}")

def maybe_opencc_to_tw(texts: List[str], cfg: TranslateConfig) -> List[str]:
    if not cfg.use_opencc:
        return texts
    try:
        import opencc
    except Exception:
        print("警告：未安裝 opencc，將跳過簡轉繁（台灣用語）步驟。", file=sys.stderr)
        return texts
    conv = opencc.OpenCC(cfg.opencc_config)
    return [conv.convert(t) for t in texts]

def write_markdown(paragraphs: List[str], out_path: str, meta: Optional[dict]=None):
    with open(out_path, "w", encoding="utf-8") as f:
        if meta:
            f.write("---\n")
            for k,v in meta.items():
                f.write(f"{k}: {v}\n")
            f.write("---\n\n")
        for p in paragraphs:
            f.write(p.strip() + "\n\n")

def write_docx(paragraphs: List[str], out_path: str):
    try:
        from docx import Document
    except Exception:
        raise RuntimeError("需要安裝 python-docx 才能輸出 .docx")
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
        doc.add_paragraph("")  # 空行
    doc.save(out_path)

def gemini_refine(msg: str) -> str:
    import os
    import requests
    gemini_api_key = os.environ.get("GEMINI_API_KEY")
    if not gemini_api_key:
        raise ValueError("請設定環境變數 GEMINI_API_KEY")
    headers = {"Authorization": f"Bearer {gemini_api_key}"}
    genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
    model = genai.GenerativeModel('gemini-2.5-flash')
    prompt = msg + '\n\n因為現在的內容可能有些生硬或不夠流暢。\n請幫我把以上文字稍微潤飾一下，使其更通順自然。注意不要有任何其他多餘的回覆'
    response = model.generate_content(prompt)
    return response.text
def main():
    ap = argparse.ArgumentParser(description="把 PDF 學術論文翻成繁體中文（含數學式保護）。")
    ap.add_argument("pdf", help="輸入 PDF 路徑")
    ap.add_argument("--out", required=True, help="輸出檔（.md 或 .docx）")
    ap.add_argument("--backend", choices=["hf", "openai", "deepl"], default="hf", help="翻譯後端（預設 hf）")
    ap.add_argument("--openai-model", default="gpt-4o-mini", help="OpenAI 模型名")
    ap.add_argument("--hf-model", default=None, help="HF 模型名（預設 Helsinki-NLP/opus-mt-en-zh；若英->中可用 Helsinki-NLP/opus-mt-en-zh）也可以用 facebook/m2m100_418M ")
    ap.add_argument("--src", dest="src_lang", default="en", help="來源語言代碼（如 en、ja、de；M2M100 需要）")
    ap.add_argument("--tgt", dest="tgt_lang", default="zh-TW", help="目標語言代碼（M2M100 支援 zh-CN/zh-TW 等）")
    ap.add_argument("--no-opencc", default=True, action="store_true", help="停用簡轉繁（台灣用語）")
    ap.add_argument("--ocr", default=True, action="store_true", help="掃描型 PDF 開啟 OCR")
    args = ap.parse_args()

    cfg = TranslateConfig(
        src_lang=args.src_lang,
        tgt_lang=args.tgt_lang,
        backend=args.backend,
        openai_model=args.openai_model,
        hf_model=args.hf_model,
        use_opencc=not args.no_opencc
    )

    # 1) 讀 PDF -> 段落
    raw_paragraphs = extract_paragraphs_from_pdf(args.pdf, use_ocr=args.ocr)
    print(f'raw_paragraphs: {raw_paragraphs}')
    # 2) 先把段落合併為適中大小批次，且對每批做數學式 mask
    batches = split_for_translation(raw_paragraphs, max_chars=600)  # 降低批次大小


    masked_batches = []
    math_maps = []
    for b in batches:
        masked, mp = mask_math(b)
        masked_batches.append(masked)
        math_maps.append(mp)

    # 3) 翻譯
    backend = build_backend(cfg)
    translated_batches = backend.translate_list(masked_batches, cfg)

    # 4) 還原數學式
    restored = [unmask_math(tb, mp) for tb, mp in zip(translated_batches, math_maps)]

    # 5) opencc 簡轉繁（若使用 OpenAI/DeepL 輸出常為簡體，可轉成台灣用語）
    restored = maybe_opencc_to_tw(restored, cfg)

    # 6) 拆回段落（以原先的空行切分）
    out_paragraphs = []
    for b in restored:
        out_paragraphs.extend([p.strip() for p in re.split(r"\n\s*\n", b) if p.strip()])

    # 7) 寫出
    meta = {"source_pdf": os.path.abspath(args.pdf)}
    if args.out.lower().endswith(".md"):
        write_markdown(out_paragraphs, args.out, meta=meta)
    elif args.out.lower().endswith(".docx"):
        write_docx(out_paragraphs, args.out)
    else:
        raise ValueError("輸出副檔名需為 .md 或 .docx")



    with open(args.out, "r", encoding="utf-8") as f:
        data = f.read()

    data = gemini_refine(data)

    with open(args.out, "w", encoding="utf-8") as f:
        f.write(data)

    print(f"✅ 完成：{args.out}")

if __name__ == "__main__":
    main()
